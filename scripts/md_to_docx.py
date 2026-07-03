"""Minimal, document-tailored Markdown -> .docx converter for DEMO-SCRIPT.md.

Handles the constructs actually used in the file: H1/H2/H3, GitHub-style tables,
unordered lists (with nesting by indent), blockquotes, horizontal rules, and
inline **bold** / `code`. Not a general Markdown engine.
"""

import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Render **bold**, `code` and plain runs into a paragraph."""
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def main(src, dst):
    with open(src, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            p = doc.add_paragraph()
            p.add_run("─" * 30).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # headings
        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading("", level=min(level, 4))
            add_inline(heading, m.group(2))
            i += 1
            continue

        # table block (line containing | then a separator row)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for idx, htext in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], htext)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx, ctext in enumerate(row):
                    if idx < len(cells):
                        cells[idx].paragraphs[0].text = ""
                        add_inline(cells[idx].paragraphs[0], ctext)
            doc.add_paragraph()
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # list item (supports one level of nesting by leading spaces)
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group(2))
            i += 1
            continue

        # ordered list item
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(dst)
    print(f"Wrote {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
